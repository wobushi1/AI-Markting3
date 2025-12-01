import sys
import json
import base64
import os
import shutil
import tempfile
from io import BytesIO
from typing import List

# PyQt5 模块
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QTextEdit, QLabel, 
                             QFileDialog, QListWidget, QSplitter, QProgressBar,
                             QLineEdit, QFormLayout, QMessageBox, QTabWidget,
                             QSizePolicy, QGroupBox, QProgressDialog)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor

# 图像处理模块
from PIL import Image
from pdf2image import convert_from_path
from openai import OpenAI

# Word 操作模块
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ==========================================
# 工具函数：获取 Poppler 路径 (新增)
# ==========================================
def get_poppler_path():
    """
    确定 Poppler 的 bin 路径。
    1. 如果是 PyInstaller 打包后的环境 (sys._MEIPASS)，返回打包内的路径。
    2. 如果是本地开发环境，返回 None (依赖系统 PATH 环境变量)。
    """
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller 将资源解压到的临时文件夹
        # 对应 spec/cmd 中的 --add-data "poppler/bin;poppler/bin"
        return os.path.join(sys._MEIPASS, 'poppler', 'bin')
    return None

# ==========================================
# 评分标准配置 (保持不变)
# ==========================================
RUBRIC_PROMPT = """
你是一位资深的高考英语阅卷专家。请对上传的手写英语作文图片进行识别、分类、评分，并提供极度详细的逐句修改意见。
注意：图片可能包含试卷的题干或表格，请只提取并批改学生手写的作文部分。

请严格按照以下 JSON 格式返回：
{
    "recognized_text": "识别出的原文...",
    "essay_type": "应用文/读后续写",
    "scores": {
        "dim1_score": 4, 
        "dim2_score": 3,
        "dim3_score": 4,
        "total": 11
    },
    "feedback_detail": {
        "content": {
            "weakness": "...",
            "suggestion": "..."
        },
        "language": {
            "sentence_corrections": [
                {
                    "original": "Original sentence...",
                    "revised": "Revised sentence...",
                    "explanation": "Grammar point..."
                }
            ],
            "general_comment": "..."
        },
        "structure": "...",
        "overall_summary": "..."
    },
    "revised_version": "Full revised essay..."
}
"""

# ==========================================
# 后端工作线程 (保持不变)
# ==========================================
class Worker(QThread):
    finished = pyqtSignal(dict, str)
    error = pyqtSignal(str, str)

    def __init__(self, file_path, api_key, model_endpoint):
        super().__init__()
        self.file_path = file_path
        self.api_key = api_key
        self.model_endpoint = model_endpoint 

    def encode_image(self, image_path):
        try:
            img = Image.open(image_path)
            if img is None: raise Exception("无法加载文件")
            if img.mode in ("RGBA", "P"): img = img.convert("RGB")
            
            max_size = 2048
            if max(img.size) > max_size:
                img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)

            buffered = BytesIO()
            img.save(buffered, format="JPEG", quality=85)
            return base64.b64encode(buffered.getvalue()).decode('utf-8')
        except Exception as e:
            raise Exception(f"文件预处理失败: {str(e)}")

    def run(self):
        try:
            base64_image = self.encode_image(self.file_path)
            # 注意：这里的 base_url 是针对火山引擎的，请确保正确
            client = OpenAI(api_key=self.api_key, base_url="https://ark.cn-beijing.volces.com/api/v3")
            
            response = client.chat.completions.create(
                model=self.model_endpoint,
                messages=[
                    {"role": "system", "content": RUBRIC_PROMPT},
                    {"role": "user", "content": [{"type": "text", "text": "批改此作文并返回JSON。"}, 
                                                 {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}]}
                ],
                temperature=0.2
            )
            content = response.choices[0].message.content.replace("```json", "").replace("```", "").strip()
            if content.startswith("json"): content = content[4:]
            
            result_json = json.loads(content)
            self.finished.emit(result_json, self.file_path)
        except Exception as e:
            self.error.emit(str(e), self.file_path)

# ==========================================
# 前端 GUI
# ==========================================
class EssayGraderApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("英语作文智能批改系统（请接入豆包模型）")
        self.resize(1400, 900)
        
        self.results_store = {} 
        self.temp_dir = tempfile.mkdtemp(prefix="essay_grader_")
        
        # --- 新增标志位：是否请求停止 ---
        self.stop_requested = False 
        
        self.init_ui()

    def closeEvent(self, event):
        try:
            shutil.rmtree(self.temp_dir)
        except:
            pass
        event.accept()

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        layout = QVBoxLayout(main_widget)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)

        # 1. API 设置
        config_group = QGroupBox("API 设置")
        config_layout = QFormLayout()
        config_layout.setContentsMargins(10, 10, 10, 10)
        
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("火山引擎 API Key")
        self.api_key_input.setEchoMode(QLineEdit.Password)
        
        self.endpoint_input = QLineEdit()
        self.endpoint_input.setPlaceholderText("接入点 ID (如 ep-2024... Vision版)")
        
        config_layout.addRow("API Key:", self.api_key_input)
        config_layout.addRow("Endpoint ID:", self.endpoint_input)
        config_group.setLayout(config_layout)
        config_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        layout.addWidget(config_group)

        # 2. 中间区域
        splitter = QSplitter(Qt.Horizontal)
        splitter.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        
        # --- 左侧控制区 ---
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        
        # 按钮区 1: 核心控制
        btn_layout_top = QHBoxLayout()
        self.btn_add = QPushButton("📂 添加文件")
        self.btn_add.clicked.connect(self.add_files)
        
        self.btn_run = QPushButton("▶ 开始批改")
        self.btn_run.clicked.connect(self.start_grading)
        self.btn_run.setStyleSheet("background-color: #007AFF; color: white; font-weight: bold;")
        
        # --- 修改点：增加停止按钮 ---
        self.btn_stop = QPushButton("🛑 停止")
        self.btn_stop.clicked.connect(self.stop_grading_process)
        self.btn_stop.setStyleSheet("background-color: #E53935; color: white; font-weight: bold;")
        self.btn_stop.setEnabled(False) # 初始禁用
        
        btn_layout_top.addWidget(self.btn_add)
        btn_layout_top.addWidget(self.btn_run)
        btn_layout_top.addWidget(self.btn_stop)
        
        # 按钮区 2: 列表管理
        btn_layout_mid = QHBoxLayout()
        self.btn_del = QPushButton("删除选中")
        self.btn_del.clicked.connect(self.delete_selected)
        
        self.btn_clear = QPushButton("清空列表")
        self.btn_clear.clicked.connect(self.clear_all)
        
        btn_layout_mid.addWidget(self.btn_del)
        btn_layout_mid.addWidget(self.btn_clear)

        self.btn_export = QPushButton("📄 导出Word文档")
        self.btn_export.clicked.connect(self.export_to_word)
        self.btn_export.setStyleSheet("background-color: #FF9800; color: white; font-weight: bold; padding: 5px;")
        self.btn_export.setEnabled(False)

        left_layout.addLayout(btn_layout_top)
        left_layout.addLayout(btn_layout_mid)
        left_layout.addWidget(self.btn_export)
        
        left_layout.addWidget(QLabel("文件列表 (支持多页PDF):"))
        self.file_list = QListWidget()
        self.file_list.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.file_list.itemClicked.connect(self.load_selected_result)
        left_layout.addWidget(self.file_list)
        
        # --- 右侧显示区 ---
        right_widget = QTabWidget()
        right_widget.setStyleSheet("QTextEdit { font-size: 16px; line-height: 1.6; }")
        
        self.text_original = QTextEdit()
        self.text_original.setReadOnly(True)
        right_widget.addTab(self.text_original, "📝 识别原文")
        
        self.text_feedback = QTextEdit()
        self.text_feedback.setReadOnly(True)
        right_widget.addTab(self.text_feedback, "📊 深度精批")
        
        self.text_revised = QTextEdit()
        self.text_revised.setReadOnly(True)
        right_widget.addTab(self.text_revised, "✨ 满分范文")

        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 4)
        layout.addWidget(splitter)

        # 3. 状态栏
        self.progress_bar = QProgressBar()
        self.status_label = QLabel("就绪")
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)

        self.setFont(QFont("Microsoft YaHei", 10))

    # ==========================
    # 逻辑功能
    # ==========================
    
    def add_files(self):
        filters = "All Support (*.png *.jpg *.jpeg *.pdf);;Images (*.png *.jpg);;PDF (*.pdf)"
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件", "", filters)
        
        if not files: return

        progress = QProgressDialog("正在解析文件...", "取消", 0, len(files), self)
        progress.setWindowModality(Qt.WindowModal)
        
        count = 0
        for f in files:
            if progress.wasCanceled(): break
            
            ext = os.path.splitext(f)[1].lower()
            filename = os.path.basename(f)
            
            if ext == '.pdf':
                try:
                    self.status_label.setText(f"正在拆分 PDF: {filename}...")
                    QApplication.processEvents()
                    
                    # === 修改点：使用 get_poppler_path() 传入正确的路径 ===
                    poppler_bin = get_poppler_path()
                    pages = convert_from_path(f, poppler_path=poppler_bin)
                    
                    for i, page in enumerate(pages):
                        page_filename = f"{os.path.splitext(filename)[0]}_Page_{i+1}.jpg"
                        temp_path = os.path.join(self.temp_dir, page_filename)
                        page.save(temp_path, "JPEG")
                        display_name = f"[PDF P{i+1}] {filename}"
                        self.add_item_to_list(display_name, temp_path)
                except Exception as e:
                    QMessageBox.warning(self, "转换失败", f"无法解析 PDF {filename}:\n请确保 Poppler 已安装。\n错误信息: {str(e)}")
            else:
                self.add_item_to_list(filename, f)
            
            count += 1
            progress.setValue(count)
        
        self.status_label.setText(f"添加完成")

    def add_item_to_list(self, display_name, file_path):
        items = [self.file_list.item(x).data(Qt.UserRole) for x in range(self.file_list.count())]
        if file_path not in items:
            list_item = self.file_list.addItem(display_name)
            self.file_list.item(self.file_list.count()-1).setData(Qt.UserRole, file_path)

    def delete_selected(self):
        row = self.file_list.currentRow()
        if row >= 0:
            item = self.file_list.takeItem(row)
            file_path = item.data(Qt.UserRole)
            if file_path in self.results_store:
                del self.results_store[file_path]
            self.refresh_ui_state()

    def clear_all(self):
        if self.file_list.count() > 0:
            reply = QMessageBox.question(self, '确认', '确定要清空所有文件和结果吗？', 
                                         QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.file_list.clear()
                self.results_store.clear()
                self.refresh_ui_state()
                self.progress_bar.setValue(0)
                self.status_label.setText("就绪")

    def refresh_ui_state(self):
        """辅助函数：清理显示区"""
        if self.file_list.count() == 0:
            self.text_original.clear()
            self.text_feedback.clear()
            self.text_revised.clear()
            self.btn_export.setEnabled(False)

    # ==========================
    # 核心：批改控制流程
    # ==========================

    def start_grading(self):
        if self.file_list.count() == 0: return
        api_key = self.api_key_input.text().strip()
        endpoint = self.endpoint_input.text().strip()
        if not api_key or not endpoint:
            QMessageBox.warning(self, "提示", "请填写API Key和Endpoint")
            return
        
        # 1. 锁定UI，重置标志位
        self.stop_requested = False
        self.btn_run.setEnabled(False)
        self.btn_stop.setEnabled(True) # 启用停止按钮
        self.btn_export.setEnabled(False)
        self.btn_del.setEnabled(False)
        self.btn_clear.setEnabled(False)
        self.btn_add.setEnabled(False)
        
        # 从头开始，或者继续未完成的？这里简化逻辑：总是从头寻找第一个未批改的
        self.process_next_file(0, api_key, endpoint)

    # --- 新增功能：停止按钮槽函数 ---
    def stop_grading_process(self):
        self.stop_requested = True
        self.btn_stop.setEnabled(False) # 防止重复点击
        self.status_label.setText("⚠️ 正在停止... 当前任务完成后将中止")

    def process_next_file(self, index, api_key, endpoint):
        # 1. 检查是否越界
        if index >= self.file_list.count():
            self.finish_grading_session("所有文件批改完成")
            return

        item = self.file_list.item(index)
        file_path = item.data(Qt.UserRole)
        display_name = item.text()
        
        # 2. 如果已经批改过，跳过
        if file_path in self.results_store:
            self.process_next_file(index + 1, api_key, endpoint)
            return

        # 3. 开始处理当前文件
        self.file_list.setCurrentRow(index)
        self.status_label.setText(f"正在处理: {display_name}")
        self.progress_bar.setValue(int((index / self.file_list.count()) * 100))

        self.worker = Worker(file_path, api_key, endpoint)
        self.worker.finished.connect(lambda res, path: self.on_result(res, path, index, api_key, endpoint))
        self.worker.error.connect(lambda err, path: self.on_error(err, path, index, api_key, endpoint))
        self.worker.start()

    def on_result(self, result, file_path, index, api_key, endpoint):
        # 保存结果
        self.results_store[file_path] = result
        item = self.file_list.item(index)
        original_text = item.text()
        if not original_text.startswith("✅"):
            item.setText(f"✅ {original_text}")
        item.setForeground(QColor("green"))
        self.display_result(result)
        
        # --- 关键修改：检查停止标志 ---
        if self.stop_requested:
            self.finish_grading_session(f"已停止。已批改 {len(self.results_store)} 份文件。")
        else:
            self.process_next_file(index + 1, api_key, endpoint)

    def on_error(self, err, file_path, index, api_key, endpoint):
        self.status_label.setText(f"错误: {err}")
        item = self.file_list.item(index)
        original_text = item.text()
        if not original_text.startswith("❌"):
            item.setText(f"❌ {original_text}")
        item.setForeground(QColor("red"))
        
        # --- 关键修改：检查停止标志 ---
        if self.stop_requested:
            self.finish_grading_session("已停止（发生错误后中断）。")
        else:
            self.process_next_file(index + 1, api_key, endpoint)

    def finish_grading_session(self, message):
        """批改会话结束（无论是完成还是停止）后的清理工作"""
        self.status_label.setText(message)
        self.progress_bar.setValue(100) if "完成" in message else None
        
        # 恢复按钮状态
        self.btn_run.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.btn_export.setEnabled(True) # 允许导出已有结果
        self.btn_del.setEnabled(True)
        self.btn_clear.setEnabled(True)
        self.btn_add.setEnabled(True)
        
        QMessageBox.information(self, "状态", f"{message}\n您可以导出当前已有的结果。")

    # ==========================
    # 显示与导出 (保持不变)
    # ==========================
    def load_selected_result(self, item):
        file_path = item.data(Qt.UserRole)
        if file_path in self.results_store:
            self.display_result(self.results_store[file_path])
        else:
            self.text_original.setText("等待处理或处理失败...")
            self.text_feedback.clear()
            self.text_revised.clear()

    def display_result(self, data):
        self.text_original.setText(f"【类型】：{data.get('essay_type', '未分类')}\n\n{data.get('recognized_text', '')}")
        self.text_revised.setText(data.get('revised_version', '暂无'))
        
        scores = data.get('scores', {})
        fb = data.get('feedback_detail', {})
        content_fb = fb.get('content', {})
        lang_fb = fb.get('language', {})
        
        html = f"""
        <h2 style='color:#333'>总分：<span style='color:#E53935; font-size:24px'>{scores.get('total', 0)}/15</span></h2>
        <table border='1' cellpadding='6' cellspacing='0' style='border-collapse:collapse; width:100%; border-color:#ddd;'>
            <tr style='background-color:#f5f5f5'>
                <th width='33%'>内容要点</th><th width='33%'>语言表达</th><th width='33%'>结构衔接</th>
            </tr>
            <tr>
                <td align='center'>{scores.get('dim1_score', 0)}/5</td>
                <td align='center'>{scores.get('dim2_score', 0)}/5</td>
                <td align='center'>{scores.get('dim3_score', 0)}/5</td>
            </tr>
        </table>
        <h3 style='background-color:#E3F2FD'>一、内容要点</h3>
        <ul>
            <li><b>🔻 不足：</b> {content_fb.get('weakness', '无')}</li>
            <li><b>💡 建议：</b> {content_fb.get('suggestion', '无')}</li>
        </ul>
        <h3 style='background-color:#FFF3E0'>二、语言表达 (逐句精改)</h3>
        """
        corrections = lang_fb.get('sentence_corrections', [])
        if corrections:
            for idx, item in enumerate(corrections, 1):
                html += f"""
                <div style='margin-bottom:10px; border-bottom:1px dashed #ccc; padding-bottom:5px;'>
                    <p><b>{idx}. 🔴 原句：</b> <span style='color:#555'>{item.get('original')}</span></p>
                    <p><b>🟢 修改：</b> <span style='color:#2E7D32; font-weight:bold'>{item.get('revised')}</span></p>
                    <p><b>📘 解析：</b> <span style='color:#1565C0'>{item.get('explanation')}</span></p>
                </div>
                """
        else:
            html += "<p>暂无具体修改建议。</p>"
            
        html += f"""
        <h3 style='background-color:#E8F5E9'>三、结构与衔接</h3>
        <p>{fb.get('structure', '无')}</p>
        <hr>
        <p><b>🌟 整体总结：</b> {fb.get('overall_summary', '')}</p>
        """
        self.text_feedback.setHtml(html)

    def export_to_word(self):
        if not self.results_store:
            QMessageBox.warning(self, "提示", "没有可导出的数据")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "保存Word文档", "批量批改结果.docx", "Word Files (*.docx)")
        if not save_path:
            return

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman' 
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei') 
        
        for i in range(self.file_list.count()):
            item = self.file_list.item(i)
            file_path = item.data(Qt.UserRole)
            display_name = item.text().replace("✅ ", "").replace("❌ ", "")
            
            if file_path not in self.results_store:
                continue
                
            data = self.results_store[file_path]
            
            doc.add_heading(f"文件：{display_name}", level=1)
            
            doc.add_heading("OCR 识别原文", level=2)
            p = doc.add_paragraph(data.get('recognized_text', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            doc.add_heading("评分详情", level=2)
            scores = data.get('scores', {})
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '维度'
            hdr_cells[1].text = '内容要点'
            hdr_cells[2].text = '语言表达'
            hdr_cells[3].text = '结构衔接'
            
            row_cells = table.rows[1].cells
            row_cells[0].text = '得分'
            row_cells[1].text = str(scores.get('dim1_score', 0))
            row_cells[2].text = str(scores.get('dim2_score', 0))
            row_cells[3].text = str(scores.get('dim3_score', 0))
            
            total_p = doc.add_paragraph()
            run = total_p.add_run(f"总分：{scores.get('total')}/15")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(14)

            fb = data.get('feedback_detail', {})
            doc.add_heading("一、内容要点", level=3)
            content_fb = fb.get('content', {})
            doc.add_paragraph(f"不足：{content_fb.get('weakness', '无')}", style='List Bullet')
            doc.add_paragraph(f"建议：{content_fb.get('suggestion', '无')}", style='List Bullet')
            
            doc.add_heading("二、语言表达与逐句修改", level=3)
            lang_fb = fb.get('language', {})
            corrections = lang_fb.get('sentence_corrections', []) if lang_fb else []
            
            if corrections:
                for idx, cor in enumerate(corrections, 1):
                    p_group = doc.add_paragraph()
                    p_group.add_run(f"{idx}. 原句：").bold = True
                    p_group.add_run(cor.get('original', '')).font.color.rgb = RGBColor(100, 100, 100)
                    
                    p_group = doc.add_paragraph()
                    p_group.add_run(f"   修改：").bold = True
                    run_revised = p_group.add_run(cor.get('revised', ''))
                    run_revised.font.color.rgb = RGBColor(0, 128, 0)
                    run_revised.bold = True
                    
                    p_group = doc.add_paragraph()
                    p_group.add_run(f"   解析：").bold = True
                    p_group.add_run(cor.get('explanation', '')).font.color.rgb = RGBColor(0, 0, 255)
                    doc.add_paragraph("") 
            else:
                doc.add_paragraph("暂无具体的逐句修改建议。")

            doc.add_heading("三、结构与整体总结", level=3)
            doc.add_paragraph(f"结构评价：{fb.get('structure', '无')}")
            doc.add_paragraph(f"整体总结：{fb.get('overall_summary', '无')}")

            doc.add_heading("满分范文参考", level=2)
            doc.add_paragraph(data.get('revised_version', '暂无'))
            doc.add_page_break()

        try:
            doc.save(save_path)
            QMessageBox.information(self, "成功", f"文档已保存至：\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败：{str(e)}\n请检查文件是否被占用。")
    
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = EssayGraderApp()
    window.show()
    sys.exit(app.exec_())
